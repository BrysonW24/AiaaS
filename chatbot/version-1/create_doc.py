pip install Document;

from docx import Document

doc = Document()
doc.add_heading("AI Strategy for MSPs", level=1)
doc.add_paragraph("This document outlines key AI use cases for Managed Service Providers.")
doc.save("your_doc.docx")
