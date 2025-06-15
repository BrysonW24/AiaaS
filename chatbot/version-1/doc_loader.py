from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

def load_docx_chunks(file_path):
    doc = Document(file_path)
    full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=50
    )
    chunks = splitter.split_text(full_text)
    return chunks
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

def load_docx_chunks(file_path):
    doc = Document(file_path)
    full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=50
    )
    chunks = splitter.split_text(full_text)
    return chunks

